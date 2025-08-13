import os
import sys
import json
import argparse
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, Any, Optional, List
from dotenv import load_dotenv
import openai
from openai import OpenAI
from prompts.collective.test_plan_generation_system_prompt import TEST_PLAN_SYSTEM_PROMPT

# Load environment variables
load_dotenv()

def generate_test_plan(api_key: Optional[str] = None, feature_info: Optional[Dict] = None, verbose: bool = False) -> tuple[bool, Dict[str, Any], str]:
    """
    Generate a test plan for a specified feature in JSON format.

    Args:
        feature_name (str): The name of the feature to create a test plan for
        api_key (str): Not used for Gemini
        feature_info (dict): Optional feature information to enhance the prompt
        verbose (bool): Whether to print verbose output

    Returns:
        tuple: (success: bool, test_plan: Dict[str, Any], raw_response: str)
               success: True if generation was successful, False otherwise
               test_plan: Generated test plan in JSON format (empty dict if failed)
               raw_response: Raw response text for debugging (empty string if failed)
    """
    try:
        openai.api_key = api_key or os.getenv("OPENAI_API_KEY")
        if not openai.api_key:
            print("Error: OpenAI API key not found. Provide it as an argument or set the OPENAI_API_KEY environment variable.")
            return False, {}, ""

        # Use Intel's internal API if needed, otherwise use standard OpenAI endpoint
        base_url = "https://apis-internal.intel.com/generativeaiinference/v4"
        client = OpenAI(api_key=openai.api_key, base_url=base_url)

        base_prompt = TEST_PLAN_SYSTEM_PROMPT
        if feature_info:
            feature_info_str = json.dumps(feature_info, indent=2)
            base_prompt += f"\n\nConsider the feature information while generating the test plan:\n {feature_info_str}"
            print(f"Generating test plan for feature: {feature_info['name']}")

        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a testing expert that creates reliable JSON-formatted test plans. \n Do not stop output until the entire JSON object is fully completed. \n Ensure every opening { or [ has a matching closing } or ]. \n Do not output explanatory text or partial data. \n If the JSON is large, always generate the full object in one response. \n If truncation is likely, indicate it clearly with a comment at the end like // TRUNCATED. \n All tests for a feature or collective must go into one file."},
                    {"role": "user", "content": base_prompt}
                ],
                response_format={"type": "json_object"},
                temperature=0.3,        # Lower temperature for more consistent formatting
                max_completion_tokens=5000
            )

            # Get the raw text response
            response_text = response.choices[0].message.content
            try:
                test_plan = json.loads(response_text)
            except json.JSONDecodeError:
                print(f"Error generating test plan in json format: {str(e)}")
                return False, {}, ""

            if verbose:
                print("Successfully generated test plan")

            return True, test_plan, response_text

        except Exception as e:
            print(f"Error during API call or JSON parsing: {str(e)}")
            return False, {}, ""

    except Exception as e:
        print(f"Error generating test plan: {str(e)}")
        return False, {}, ""

def create_test_plan_document(test_plan: Dict[str, Any], output_file: str, feature_info) -> bool:
    """Create a Word document from the test plan.

    Args:
        test_plan: The test plan dictionary
        output_file: Path to save the DOCX file
        feature_info: Feature information (dict or string)

    Returns:
        bool: True if document was created successfully, False otherwise
    """
    try:
        doc = Document()

        # Add title - handle both feature_info as dict or string
        if isinstance(feature_info, dict) and 'name' in feature_info:
            title_text = f'Test Plan: {feature_info["name"]}'
        elif isinstance(feature_info, str):
            title_text = f'Test Plan: {feature_info}'
        elif 'test_plan' in test_plan:
            title_text = f'Test Plan: {test_plan["test_plan"]}'
        elif 'test_category' in test_plan:
            title_text = f'Test Plan: {test_plan["test_category"]}'
        else:
            title_text = 'Test Plan'

        title = doc.add_heading(title_text, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Handle the new JSON structure with "tests" containing multiple categories
        if 'tests' in test_plan:
            tests_data = test_plan['tests']

            # Handle case where tests is a dict with multiple test categories
            if isinstance(tests_data, dict):
                for key, category_data in tests_data.items():
                    if isinstance(category_data, dict) and 'test_category' in category_data:
                        # Add category heading
                        cat_heading = doc.add_heading(f'Category: {category_data["test_category"]}', level=1)

                        # Add implementation file if present
                        if 'implementation_file' in category_data:
                            doc.add_paragraph(f'Implementation File: {category_data["implementation_file"]}')
                            doc.add_paragraph()  # Add spacing

                        # Add test cases
                        if 'test_cases' in category_data:
                            test_cases = category_data['test_cases']
                            if isinstance(test_cases, list):
                                for test_case in test_cases:
                                    if isinstance(test_case, dict):
                                        tc_heading = doc.add_heading(f'Test Case: {test_case.get("test_id", "Unknown")}', level=2)
                                        doc.add_paragraph(f'Description: {test_case.get("description", "No description provided")}')
                                        doc.add_paragraph()  # Add spacing
                            elif isinstance(test_cases, dict):
                                # Handle case where test_cases is a dict
                                for test_id, test_info in test_cases.items():
                                    if isinstance(test_info, dict):
                                        tc_heading = doc.add_heading(f'Test Case: {test_id}', level=2)
                                        doc.add_paragraph(f'Description: {test_info.get("description", "No description provided")}')
                                    else:
                                        tc_heading = doc.add_heading(f'Test Case: {test_id}', level=2)
                                        doc.add_paragraph(f'Description: {test_info}')
                                    doc.add_paragraph()  # Add spacing

            # Handle case where tests is a list
            elif isinstance(tests_data, list):
                for category_data in tests_data:
                    if isinstance(category_data, dict) and 'test_category' in category_data:
                        cat_heading = doc.add_heading(f'Category: {category_data["test_category"]}', level=1)

                        if 'implementation_file' in category_data:
                            doc.add_paragraph(f'Implementation File: {category_data["implementation_file"]}')
                            doc.add_paragraph()

                        if 'test_cases' in category_data:
                            for test_case in category_data['test_cases']:
                                tc_heading = doc.add_heading(f'Test Case: {test_case.get("test_id", "Unknown")}', level=2)
                                doc.add_paragraph(f'Description: {test_case.get("description", "No description provided")}')
                                doc.add_paragraph()

        # Handle legacy structure (backward compatibility)
        elif 'test_category' in test_plan and 'test_cases' in test_plan:
            cat_heading = doc.add_heading(f'Category: {test_plan["test_category"]}', level=1)
            for test_case in test_plan['test_cases']:
                tc_heading = doc.add_heading(f'Test Case {test_case["test_id"]}', level=2)
                doc.add_paragraph(f'Description: {test_case["description"]}')
                doc.add_paragraph()

        # If no recognizable structure, just dump the content
        else:
            doc.add_paragraph("Test Plan Content:")
            doc.add_paragraph(json.dumps(test_plan, indent=2))

        doc.save(output_file)
        return True

    except Exception as e:
        print(f"Error creating Word document: {e}")
        import traceback
        traceback.print_exc()
        return False

def generate_test_plan_files(output_file: str, json_file: str, feature_info_file: Optional[str] = None, verbose: bool = False) -> bool:
    """
    Generate test plan files (both DOCX and JSON) from feature information.

    Args:
        output_file: Path for the DOCX output file
        json_file: Path for the JSON output file
        feature_info_file: Optional path to feature info JSON file
        verbose: Enable verbose output

    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Load feature info if provided
        feature_info = None
        if feature_info_file and os.path.exists(feature_info_file):
            try:
                with open(feature_info_file, 'r') as f:
                    feature_info = json.load(f)
            except Exception as e:
                print(f"Error loading feature info file: {e}")
                return False

        # Generate test plan
        success, test_plan, raw_response = generate_test_plan(
            feature_info=feature_info,
            verbose=verbose
        )

        if not success:
            print("Failed to generate test plan")
            return False

        # Save test plan as JSON
        try:
            with open(json_file, 'w') as f:
                json.dump(test_plan, f, indent=2)
        except Exception as e:
            print(f"Error saving JSON file: {e}")
            return False

        # Create Word document
        if not create_test_plan_document(test_plan, output_file, feature_info):
            print("Failed to create Word document")
            return False

        if verbose:
            print(f"Test plan saved to: {output_file}")
            print(f"JSON test plan saved to: {json_file}")

        return True

    except Exception as e:
        print(f"Error generating test plan files: {e}")
        import traceback
        traceback.print_exc()
        return False
